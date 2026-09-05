import sys
import unittest
import tempfile
import json
from zipfile import ZipFile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "document_engine"))
from template_security import validate_template_syntax
from analyze_template import analyze_template
from generate_document import generate_document
from docx import Document

class TemplateSecurityTests(unittest.TestCase):
    def test_valid_placeholder(self):
        self.assertEqual(validate_template_syntax("{{ member.fullName }}"), ["member.fullName"])

    def test_rejects_jinja_block(self):
        with self.assertRaises(ValueError):
            validate_template_syntax("{% for x in member %}{{ x }}{% endfor %}")

    def test_rejects_filter_and_call(self):
        with self.assertRaises(ValueError):
            validate_template_syntax("{{ member.fullName|upper }} {{ cycler.__init__() }}")

    def test_analyzer_accepts_valid_docx_placeholder(self):
        with tempfile.TemporaryDirectory() as directory:
            target = Path(directory) / "valid.docx"
            with ZipFile(target, "w") as archive:
                archive.writestr("[Content_Types].xml", "<Types/>")
                archive.writestr("word/document.xml", "<w:document>{{ member.fullName }}</w:document>")
            result = analyze_template(str(target))
            self.assertTrue(result["ready"])
            self.assertEqual(result["fields"], ["member.fullName"])

    def test_analyzer_rejects_jinja_in_docx(self):
        with tempfile.TemporaryDirectory() as directory:
            target = Path(directory) / "unsafe.docx"
            with ZipFile(target, "w") as archive:
                archive.writestr("[Content_Types].xml", "<Types/>")
                archive.writestr("word/document.xml", "<w:document>{% for x in member %}</w:document>")
            with self.assertRaises(ValueError):
                analyze_template(str(target))

    def test_generator_renders_allowlisted_placeholder(self):
        with tempfile.TemporaryDirectory() as directory:
            template = Path(directory) / "template.docx"
            output = Path(directory) / "generated.docx"
            data = Path(directory) / "data.json"
            document = Document()
            document.add_paragraph("{{ member.fullName }}")
            document.save(template)
            data.write_text(json.dumps({"member": {"fullName": "Pessoa Teste"}}), encoding="utf-8")
            generate_document(str(template), str(output), str(data))
            rendered = Document(output)
            self.assertIn("Pessoa Teste", "\n".join(p.text for p in rendered.paragraphs))

    def test_document_types_and_xml_escaping(self):
        for title, with_representative in [("Termo", True), ("Declaração", False)]:
            with self.subTest(title=title), tempfile.TemporaryDirectory() as directory:
                template = Path(directory) / "template.docx"
                output = Path(directory) / "output.docx"
                data = Path(directory) / "data.json"
                source = Document()
                source.add_paragraph(title + " {{ member.fullName }}")
                context = {"member": {"fullName": "Ana & João <teste>"}}
                if with_representative:
                    source.add_paragraph("{{ representative.fullName }}")
                    context["representative"] = {"fullName": "Representante Teste"}
                source.save(template)
                data.write_text(json.dumps(context), encoding="utf-8")
                generate_document(str(template), str(output), str(data))
                rendered = "\n".join(p.text for p in Document(output).paragraphs)
                self.assertIn("Ana & João <teste>", rendered)
                if with_representative:
                    self.assertIn("Representante Teste", rendered)

if __name__ == "__main__":
    unittest.main()
